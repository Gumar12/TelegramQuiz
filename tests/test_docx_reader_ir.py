import base64
import unicodedata

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK

from backend.parsers.docx_reader import read_docx_to_ir


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def _save_docx(doc: Document, tmp_path, name: str = "source.docx"):
    path = tmp_path / name
    doc.save(path)
    return path


def _write_png(tmp_path, name: str = "image.png"):
    path = tmp_path / name
    path.write_bytes(PNG_1X1)
    return path


def test_paragraphs_headings_and_blank_before(tmp_path):
    doc = Document()
    doc.add_paragraph("Обычный абзац")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_heading("Тимур", level=1)

    ir = read_docx_to_ir(_save_docx(doc, tmp_path))

    assert [block.kind for block in ir.blocks] == ["paragraph", "heading"]
    assert ir.blocks[0].text == "Обычный абзац"
    assert ir.blocks[0].style == "Normal"
    assert ir.blocks[0].blank_before == 0
    assert ir.blocks[0].source_ref["paragraph_index"] == 0

    heading = ir.blocks[1]
    assert heading.text == "Тимур"
    assert heading.style == "Heading 1"
    assert heading.blank_before == 2
    assert heading.source_ref["paragraph_index"] == 3


def test_bold_answer_span_is_preserved(tmp_path):
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("B) ")
    answer = paragraph.add_run("Керей и Жанибек")
    answer.bold = True

    ir = read_docx_to_ir(_save_docx(doc, tmp_path))

    block = ir.blocks[0]
    assert block.kind == "paragraph"
    assert block.text == "B) Керей и Жанибек"
    assert len(block.bold_spans) == 1
    span = block.bold_spans[0]
    assert (span.start, span.end, span.text) == (3, 18, "Керей и Жанибек")
    assert span.source_ref["paragraph_index"] == 0
    assert span.source_ref["run_index"] == 1
    assert span.source_ref["block_id"] == block.block_id


def test_images_are_blocks_in_document_order_and_saved_only_to_media_dir(tmp_path):
    image_path = _write_png(tmp_path)
    media_dir = tmp_path / "media"
    doc = Document()
    doc.add_paragraph("Перед картинкой")
    doc.add_paragraph().add_run().add_picture(str(image_path))
    doc.add_paragraph("Между картинками")
    doc.add_paragraph().add_run().add_picture(str(image_path))

    ir = read_docx_to_ir(_save_docx(doc, tmp_path), media_dir=media_dir)

    assert [block.kind for block in ir.blocks] == ["paragraph", "image", "paragraph", "image"]
    assert [block.text for block in ir.blocks if block.kind == "paragraph"] == [
        "Перед картинкой",
        "Между картинками",
    ]
    image_blocks = [block for block in ir.blocks if block.kind == "image"]
    assert [path.split("/")[-1] for block in image_blocks for path in block.media_refs] == [
        "image_001.png",
        "image_002.png",
    ]
    for block in image_blocks:
        saved_path = media_dir / block.media_refs[0].split("/")[-1]
        assert saved_path.exists()
        assert saved_path.parent == media_dir
        assert block.source_ref["relationship_id"].startswith("rId")
        assert block.source_ref["run_index"] == 0


def test_images_are_not_written_when_media_dir_is_omitted(tmp_path):
    image_path = _write_png(tmp_path)
    doc = Document()
    doc.add_paragraph().add_run().add_picture(str(image_path))

    ir = read_docx_to_ir(_save_docx(doc, tmp_path))

    assert len(ir.blocks) == 1
    assert ir.blocks[0].kind == "image"
    assert ir.blocks[0].media_refs == []
    assert not (tmp_path / "image_001.png").exists()


def test_page_break_marks_next_block(tmp_path):
    doc = Document()
    doc.add_paragraph("До разрыва")
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("После разрыва")

    ir = read_docx_to_ir(_save_docx(doc, tmp_path))

    assert [block.text for block in ir.blocks] == ["До разрыва", "После разрыва"]
    assert ir.blocks[0].page_break_before is False
    assert ir.blocks[1].page_break_before is True
    assert ir.blocks[1].source_ref["page_break_before_source"]["break_kind"] == "page"


def test_section_break_marks_next_block_when_python_docx_exposes_paragraph_sectpr(tmp_path):
    doc = Document()
    doc.add_paragraph("Первая секция")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Вторая секция")

    ir = read_docx_to_ir(_save_docx(doc, tmp_path))

    assert [block.text for block in ir.blocks] == ["Первая секция", "Вторая секция"]
    assert ir.blocks[0].section_break_before is False
    assert ir.blocks[1].section_break_before is True
    assert ir.blocks[1].source_ref["section_break_before_source"]["break_kind"] == "section"


def test_unicode_text_is_nfc_and_readable(tmp_path):
    doc = Document()
    doc.add_paragraph("Cafe\u0301, Қазақ тілі, Әліпби")

    ir = read_docx_to_ir(_save_docx(doc, tmp_path))

    assert ir.blocks[0].text == "Café, Қазақ тілі, Әліпби"
    assert unicodedata.is_normalized("NFC", ir.blocks[0].text)
    assert "\\u" not in ir.blocks[0].text


def test_reader_does_not_classify_questions_or_answers(tmp_path):
    doc = Document()
    doc.add_paragraph("Кто считается основателем Казахского ханства?")
    doc.add_paragraph("A) Абылай хан")
    bold = doc.add_paragraph("B) ")
    answer = bold.add_run("Керей и Жанибек")
    answer.bold = True
    doc.add_paragraph("Ответ: B")

    ir = read_docx_to_ir(_save_docx(doc, tmp_path))

    assert [block.kind for block in ir.blocks] == [
        "paragraph",
        "paragraph",
        "paragraph",
        "paragraph",
    ]
    data = ir.to_dict()
    serialized_keys = set().union(*(block.keys() for block in data["blocks"]))
    assert "question" not in serialized_keys
    assert "answers" not in serialized_keys
    assert "correct" not in serialized_keys
    assert [block["text"] for block in data["blocks"]] == [
        "Кто считается основателем Казахского ханства?",
        "A) Абылай хан",
        "B) Керей и Жанибек",
        "Ответ: B",
    ]

def test_rendered_page_break_is_ignored_as_layout_metadata(tmp_path):
    from docx.oxml import OxmlElement

    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("До разрыва")
    paragraph._p.append(OxmlElement("w:lastRenderedPageBreak"))
    paragraph.add_run(" после разрыва")
    doc.add_paragraph("Следующий абзац")
    docx_path = tmp_path / "rendered_page_break.docx"
    doc.save(docx_path)

    ir = read_docx_to_ir(docx_path)

    assert [block.text for block in ir.blocks] == ["До разрыва после разрыва", "Следующий абзац"]
    assert all(block.page_break_before is False for block in ir.blocks)
    assert all("page_break_before_source" not in block.source_ref for block in ir.blocks)
