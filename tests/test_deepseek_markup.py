import json
from pathlib import Path

import pytest
from docx import Document

from backend import deepseek_markup
from backend import studio_api
from backend.studio_jobs import JobManager


def _patch_raw_response(monkeypatch, raw_response: dict) -> None:
    class FakeResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def read(self):
            return json.dumps(raw_response).encode("utf-8")

    monkeypatch.setattr(
        deepseek_markup.urllib.request,
        "urlopen",
        lambda *args, **kwargs: FakeResponse(),
    )


def _request_markup():
    return deepseek_markup.request_markup(
        "# DOCX_BLOCK_STREAM\n",
        api_key="token",
        model="deepseek-v4-flash",
        base_url="https://api.deepseek.com",
        timeout_seconds=1,
        max_tokens=1000,
    )


def test_blocks_markdown_from_docx_marks_options_and_bold(tmp_path: Path):
    docx_path = tmp_path / "quiz.docx"
    document = Document()
    document.add_paragraph("Контекст источника")
    document.add_paragraph("Какой ответ правильный?")
    document.add_paragraph("A) Неверно")
    bold_option = document.add_paragraph()
    run = bold_option.add_run("B) Верно")
    run.bold = True
    document.add_paragraph("C) Тоже неверно")
    document.save(docx_path)

    blocks_md = deepseek_markup.blocks_markdown_from_docx(docx_path, media_dir=tmp_path / "media")

    assert "[BLOCK p-0003 | option | bold=false" in blocks_md
    assert "label=A" in blocks_md
    assert "[BLOCK p-0004 | option | bold=true" in blocks_md
    assert "label=B" in blocks_md
    assert "B) Верно" in blocks_md


def test_request_markup_parses_deepseek_json_content(monkeypatch):
    raw_response = {
        "model": "deepseek-v4-flash",
        "choices": [
            {
                "message": {
                    "content": json.dumps(
                        {
                            "document_id": "source",
                            "questions": [
                                {
                                    "id": "q001",
                                    "question_block_ids": ["p-0001"],
                                    "option_block_ids": ["p-0002", "p-0003", "p-0004"],
                                    "correct_option_block_ids": ["p-0002"],
                                }
                            ],
                        }
                    )
                }
            }
        ],
    }

    class FakeResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def read(self):
            return json.dumps(raw_response).encode("utf-8")

    monkeypatch.setattr(deepseek_markup.urllib.request, "urlopen", lambda *args, **kwargs: FakeResponse())

    markup, raw = deepseek_markup.request_markup(
        "# DOCX_BLOCK_STREAM\n",
        api_key="token",
        model="deepseek-v4-flash",
        base_url="https://api.deepseek.com",
        timeout_seconds=1,
        max_tokens=1000,
    )

    assert raw["model"] == "deepseek-v4-flash"
    assert markup["document_id"] == "source"
    assert markup["questions"][0]["question_block_ids"] == ["p-0001"]


def test_request_markup_rejects_truncated_finish_reason_length(monkeypatch):
    raw_response = {
        "model": "deepseek-v4-flash",
        "choices": [
            {
                "finish_reason": "length",
                "message": {"content": '{"document_id": "source", "questions": ['},
            }
        ],
    }
    _patch_raw_response(monkeypatch, raw_response)

    with pytest.raises(deepseek_markup.DeepSeekMarkupError) as excinfo:
        _request_markup()
    assert "truncated" in str(excinfo.value)
    assert "finish_reason=length" in str(excinfo.value)


def test_request_markup_rejects_empty_content_without_trusting_reasoning(monkeypatch):
    raw_response = {
        "model": "deepseek-v4-flash",
        "choices": [
            {
                "finish_reason": "stop",
                "message": {
                    "content": "",
                    # reasoning_content must NOT be used as the final JSON answer.
                    "reasoning_content": '{"document_id": "source", "questions": []}',
                },
            }
        ],
    }
    _patch_raw_response(monkeypatch, raw_response)

    with pytest.raises(deepseek_markup.DeepSeekMarkupError) as excinfo:
        _request_markup()
    assert "empty message content" in str(excinfo.value)


def test_create_from_docx_ai_job_builds_final_quiz(tmp_path: Path, monkeypatch):
    blocks_md = "\n".join(
        [
            "# DOCX_BLOCK_STREAM",
            "document_id: source.docx",
            "",
            "[BLOCK p-0001 | paragraph | bold=false]",
            "Вопрос?",
            "",
            "[BLOCK p-0002 | option | label=A | bold=true]",
            "A) Верно",
            "",
            "[BLOCK p-0003 | option | label=B | bold=false]",
            "B) Неверно",
            "",
            "[BLOCK p-0004 | option | label=C | bold=false]",
            "C) Тоже неверно",
            "",
        ]
    )
    markup = {
        "document_id": "source.docx",
        "questions": [
            {
                "id": "q001",
                "question_block_ids": ["p-0001"],
                "option_block_ids": ["p-0002", "p-0003", "p-0004"],
                "correct_option_block_ids": ["p-0002"],
                "context_block_ids": [],
                "media_ids": [],
                "confidence": 0.95,
                "warnings": [],
            }
        ],
        "context_regions": [],
    }

    monkeypatch.setattr(studio_api.config, "DEEPSEEK_API_KEY", "token")
    monkeypatch.setattr(studio_api.config, "DEEPSEEK_MODEL", "deepseek-v4-flash")
    monkeypatch.setattr(studio_api.config, "DEEPSEEK_BASE_URL", "https://api.deepseek.com")
    monkeypatch.setattr(studio_api.config, "DEEPSEEK_TIMEOUT_SECONDS", 1)
    monkeypatch.setattr(studio_api.config, "DEEPSEEK_MAX_TOKENS", 1000)
    monkeypatch.setattr(
        studio_api.deepseek_markup,
        "blocks_markdown_from_docx",
        lambda *args, **kwargs: blocks_md,
    )
    monkeypatch.setattr(
        studio_api.deepseek_markup,
        "request_markup",
        lambda *args, **kwargs: (markup, {"model": "deepseek-v4-flash", "usage": {"prompt_tokens": 10, "completion_tokens": 20}}),
    )

    manager = JobManager()
    job = manager.create_job("create-from-docx-ai")
    runner = studio_api._create_from_docx_ai_job(
        docx_path=tmp_path / "source.docx",
        source_path=tmp_path / "questions_v2.json",
        media_dir=tmp_path / "media",
        output_dir=tmp_path / "quizzes",
        workspace_dir=tmp_path,
        runtime_dir=tmp_path,
        title="AI quiz",
        description="AI description",
    )

    result = runner(job.id, manager)

    assert result["groups"][0]["id"] == "AI_quiz"
    output_path = Path(result["groups"][0]["output"])
    payload = json.loads(output_path.read_text(encoding="utf-8"))
    assert payload["quiz_title"] == "AI quiz"
    assert payload["questions"][0]["question"] == "Вопрос?"
    assert payload["questions"][0]["options"] == ["Верно", "Неверно", "Тоже неверно"]
    assert payload["questions"][0]["correct"] == 1
    assert payload["ai_markup"]["provider"] == "deepseek"
    assert (tmp_path / "questions_v2.json").exists()
