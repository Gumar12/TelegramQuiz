from docx import Document

from backend.docx_to_quiz_json_v2 import (
    OPTION_MAX_CHARS,
    attach_distractors,
    build_output,
    compact_for_option,
    count_unsupported_media,
    format_group_summary,
    iter_docx_blocks,
    parse_args,
    parse_blocks_to_items,
)


def test_parse_args_supports_show_groups_flag():
    args = parse_args(
        [
            "--input",
            "source.docx",
            "--output",
            "questions.json",
            "--show-groups",
        ]
    )

    assert args.show_groups is True


def test_format_group_summary_lists_group_counts():
    data = {
        "questions": [
            {"date": "19 мая", "section": "УТРО"},
            {"date": "19 мая", "section": "УТРО"},
            {"date": "19 мая", "section": "ОБЕД"},
        ]
    }

    assert format_group_summary(data) == [
        "Групп: 2",
        "19 мая УТРО: 2",
        "19 мая ОБЕД: 1",
    ]


def test_compact_for_option_extracts_short_semantic_answers():
    assert (
        compact_for_option(
            "Кюй Курмангазы Сагырбайулы Кишкентай был создан как отклик на восстание под предводительством Исатая Тайманова."
        )
        == "Кишкентай"
    )
    assert (
        compact_for_option(
            "В этот день была принята Декларация о государственном суверенитете Казахской ССР, заложившая правовую основу."
        )
        == "Декларация о государственном суверенитете Казахской ССР"
    )
    assert (
        compact_for_option(
            "Внедрение колониальной системы управления привело к ограничению власти биев и султанов"
        )
        == "ограничение власти биев и султанов"
    )


def test_parse_blocks_resets_context_and_media_on_new_section():
    blocks = [
        {"type": "text", "text": "11 мая"},
        {"type": "text", "text": "УТРО"},
        {"type": "text", "text": "Контекст Nº2"},
        {"type": "text", "text": "(Фото Тимура и текст о нём)"},
        {"type": "image", "path": "media/image_001.png"},
        {"type": "image", "path": "media/image_002.jpg"},
        {"type": "text", "text": "На портрете изображен:"},
        {"type": "text", "text": "Эмир Тимур"},
        {
            "type": "text",
            "text": "Личности, изображенной на портрете, был присужден статус Эмир, так как он: не был чингизидом",
        },
        {"type": "text", "text": "ОБЕД"},
        {
            "type": "text",
            "text": "Идея возрождения казахского языка и культуры в 60-х годах XX века связана с организацией: Жас Тулпар",
        },
    ]

    items = parse_blocks_to_items(blocks)

    assert len(items) == 3
    assert items[0]["context_title"] == "Контекст Nº2"
    assert items[0]["media"] == ["media/image_001.png", "media/image_002.jpg"]
    assert items[0]["type"] == "media_context_quiz"
    assert items[1]["context_title"] == "Контекст Nº2"
    assert items[1]["media"] == ["media/image_001.png", "media/image_002.jpg"]
    assert items[1]["type"] == "media_context_quiz"

    assert items[2]["section"] == "ОБЕД"
    assert items[2]["context_title"] == ""
    assert items[2]["context"] == ""
    assert items[2]["media"] == []
    assert items[2]["type"] == "simple_quiz"


def test_parse_blocks_keeps_one_context_for_multiple_following_questions():
    blocks = [
        {"type": "text", "text": "11 мая"},
        {"type": "text", "text": "УТРО"},
        {"type": "text", "text": "Контекст Nº3"},
        {"type": "text", "text": "Прочитайте текст и рассмотрите портрет."},
        {"type": "image", "path": "media/portrait.jpg"},
        {"type": "text", "text": "Кто изображен на портрете?"},
        {"type": "text", "text": "Эмир Тимур"},
        {"type": "text", "text": "Год, о котором говорится в контексте:"},
        {"type": "text", "text": "1370"},
    ]

    items = parse_blocks_to_items(blocks)

    assert len(items) == 2
    assert [item["question"] for item in items] == [
        "Кто изображен на портрете?",
        "Год, о котором говорится в контексте",
    ]
    assert [item["correct_answer"] for item in items] == ["Эмир Тимур", "1370"]
    assert all(item["context_title"] == "Контекст Nº3" for item in items)
    assert all(item["context"] == "Прочитайте текст и рассмотрите портрет." for item in items)
    assert all(item["media"] == ["media/portrait.jpg"] for item in items)


def test_parse_blocks_attaches_image_between_option_quizzes_to_next_question():
    blocks = [
        {"type": "text", "text": "11 мая"},
        {"type": "text", "text": "ОБЕД"},
        {"type": "text", "text": "В 1870 году Казахстан покинули казахи во главе с султаном:"},
        {"type": "text", "text": "A) Ахмета Жантореулы"},
        {"type": "text", "text": "B) Баймагамбета Айшуакулы"},
        {"type": "text", "text": "C) Хангали Арыстанулы", "bold_text": "C) Хангали Арыстанулы"},
        {"type": "text", "text": "D) Арслана Жантореулы"},
        {"type": "image", "path": "media/map.jpg"},
        {"type": "text", "text": "Событие, представленное на карте:"},
        {"type": "text", "text": "A) новая-экономическая политика"},
        {"type": "text", "text": "B) Гражданская война"},
        {"type": "text", "text": "C) «Малый Октябрь»"},
        {"type": "text", "text": "D) индустриализация", "bold_text": "D) индустриализация"},
    ]

    items = parse_blocks_to_items(blocks)

    assert len(items) == 2
    assert items[0]["question"] == "В 1870 году Казахстан покинули казахи во главе с султаном"
    assert items[0]["media"] == []
    assert items[1]["question"] == "Событие, представленное на карте"
    assert items[1]["media"] == ["media/map.jpg"]
    assert items[1]["type"] == "media_context_quiz"


def test_attach_distractors_generates_contextual_options_for_timur_block():
    blocks = [
        {"type": "text", "text": "11 мая"},
        {"type": "text", "text": "УТРО"},
        {"type": "text", "text": "Контекст Nº2"},
        {"type": "text", "text": "Текст о Тимуре и его государстве."},
        {"type": "image", "path": "media/timur.jpg"},
        {"type": "text", "text": "На портрете изображен:"},
        {"type": "text", "text": "Эмир Тимур"},
        {"type": "text", "text": "Личность, изображенная на портрете правила в:"},
        {"type": "text", "text": "XIV веке"},
        {"type": "text", "text": "Личности, изображенной на портрете, был присужден статус Эмир, так как он:"},
        {"type": "text", "text": "не был чингизидом"},
        {"type": "text", "text": "Государство Эмира Тимура располагалось на землях потомков Чингизхана:"},
        {"type": "text", "text": "Улуса Чагатая"},
        {"type": "text", "text": "Выберите верные утверждения про Эмира Тимура:"},
        {"type": "text", "text": "Создал государство в Мавераннахре"},
    ]

    items = parse_blocks_to_items(blocks)
    portrait, century, reason, ulus, statement = items[:5]

    assert portrait["options"] == [
        "Эмир Тимур",
        "Чингисхан",
        "Абылай хан",
        "Тауке хан",
    ]
    assert century["options"] == ["XIV веке", "XIII веке", "XV веке", "XVI веке"]
    assert reason["options"] == [
        "не был чингизидом",
        "был чингизидом",
        "принял ханский титул",
        "происходил из рода Джучи",
    ]
    assert ulus["options"] == [
        "Улуса Чагатая",
        "Улуса Джучи",
        "Могулистана",
        "Золотой Орды",
    ]
    assert statement["type"] == "needs_distractor_review"
    assert statement["options"] == []
    assert statement["correct"] is None
    assert statement["distractors_source"] == "needs_contextual_distractors"
    for item in [portrait, century, reason, ulus]:
        assert all("…" not in option for option in item["options"])
        assert all(len(option) <= 100 for option in item["options"])
        assert not any("колониальной системы" in option for option in item["options"])


def test_statement_selection_without_source_options_does_not_get_chingisid_fallbacks():
    blocks = [
        {"type": "text", "text": "Контекст Nº1"},
        {
            "type": "text",
            "text": "Текст об улусе Джучи и владениях Чагатая.",
        },
        {"type": "text", "text": "Выберите верные утверждения:"},
        {"type": "text", "text": "I. В тексте описываются территориальные владения Чагатая;"},
    ]

    items = parse_blocks_to_items(blocks)

    assert len(items) == 1
    assert items[0]["question"] == "Выберите верные утверждения"
    assert items[0]["correct_answer"] == "I. В тексте описываются территориальные владения Чагатая;"
    assert items[0]["type"] == "needs_distractor_review"
    assert items[0]["options"] == []
    assert items[0]["correct"] is None


def test_roman_statement_sequence_block_uses_prompt_as_question_and_statements_as_context():
    blocks = [
        {"type": "text", "text": "Установите верную хронологическую последовательность событий:"},
        {"type": "text", "text": "I. Первая Конституция РК;"},
        {"type": "text", "text": "II. Размежевание границ республик Средней Азии;"},
        {"type": "text", "text": "III. Возвращение исконного имени «казах» народу;"},
        {"type": "text", "text": "IV. Декларация прав трудящихся."},
        {"type": "text", "text": "A) II, I, III, IV"},
        {"type": "text", "text": "B) I, IV, II, III"},
        {"type": "text", "text": "C) III, II, I, IV"},
        {"type": "text", "text": "D) IV, II, III, I", "bold_text": "D) IV, II, III, I"},
    ]

    items = parse_blocks_to_items(blocks)

    assert len(items) == 1
    assert items[0]["question"] == "Установите верную хронологическую последовательность событий"
    assert "I. Первая Конституция РК;" in items[0]["context"]
    assert "IV. Декларация прав трудящихся." in items[0]["context"]
    assert items[0]["options"] == [
        "II, I, III, IV",
        "I, IV, II, III",
        "III, II, I, IV",
        "IV, II, III, I",
    ]
    assert items[0]["correct"] == 4
    assert items[0]["correct_answer"] == "IV, II, III, I"


def test_quoted_context_with_colon_is_not_split_into_inline_question():
    blocks = [
        {
            "type": "text",
            "text": "«Телесцы хотели одного – уничтожить жужаней, и Бумынь это знал.",
        },
        {
            "type": "text",
            "text": "Разгром жужаней. Разгневанный хан ответил грубо: «Ты мой плавильщик!»",
        },
        {
            "type": "text",
            "text": "Бумынь принял титул Иль-хан, но в конце 552 г. умер».",
        },
        {"type": "text", "text": "Титул правителя тюрков:"},
        {"type": "text", "text": "A) гурхан"},
        {"type": "text", "text": "B) каган", "bold_text": "B) каган"},
        {"type": "text", "text": "C) хунтайджи"},
        {"type": "text", "text": "D) джабгу"},
    ]

    items = parse_blocks_to_items(blocks)

    assert len(items) == 1
    assert items[0]["question"] == "Титул правителя тюрков"
    assert "Разгром жужаней." in items[0]["context"]
    assert items[0]["options"] == ["гурхан", "каган", "хунтайджи", "джабгу"]
    assert items[0]["correct"] == 2


def test_unknown_items_are_marked_for_distractor_review_instead_of_getting_global_junk():
    blocks = [
        {"type": "text", "text": "11 мая"},
        {"type": "text", "text": "Неизвестный локальный вопрос:"},
        {"type": "text", "text": "локальный ответ"},
        {"type": "text", "text": "Другая несвязанная тема:"},
        {"type": "text", "text": "чужой ответ"},
    ]

    items = parse_blocks_to_items(blocks)

    assert items[0]["type"] == "needs_distractor_review"
    assert items[0]["correct_answer"] == "локальный ответ"
    assert items[0]["options"] == []
    assert items[0]["correct"] is None


def test_iter_docx_blocks_preserves_bold_text_for_ready_quiz_options(tmp_path):
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("B) ")
    answer_run = paragraph.add_run("Керей и Жанибек")
    answer_run.bold = True

    docx_path = tmp_path / "ready_quiz.docx"
    doc.save(docx_path)

    blocks = iter_docx_blocks(docx_path, tmp_path / "media")

    assert blocks == [
        {
            "type": "text",
            "text": "B) Керей и Жанибек",
            "bold_text": "Керей и Жанибек",
            "has_bold": True,
        }
    ]


def test_parse_blocks_extracts_ready_quiz_with_bold_correct_option():
    blocks = [
        {"type": "text", "text": "11 мая"},
        {"type": "text", "text": "УТРО"},
        {"type": "text", "text": "Кто считается основателем Казахского ханства?"},
        {"type": "text", "text": "A) Абылай хан"},
        {"type": "text", "text": "B) Керей и Жанибек", "bold_text": "Керей и Жанибек", "has_bold": True},
        {"type": "text", "text": "C) Кенесары хан"},
        {"type": "text", "text": "D) Тауке хан"},
    ]

    items = parse_blocks_to_items(blocks)

    assert len(items) == 1
    assert items[0]["question"] == "Кто считается основателем Казахского ханства?"
    assert items[0]["options"] == ["Абылай хан", "Керей и Жанибек", "Кенесары хан", "Тауке хан"]
    assert items[0]["correct"] == 2
    assert items[0]["correct_answer"] == "Керей и Жанибек"
    assert items[0]["type"] == "multiple_choice"
    assert items[0]["distractors_source"] == "source_document_bold"


def test_parse_blocks_extracts_ready_quiz_with_answer_letter_line():
    blocks = [
        {"type": "text", "text": "Capital of Kazakhstan:"},
        {"type": "text", "text": "A) Almaty\nB) Astana\nC) Shymkent\nD) Aktobe"},
        {"type": "text", "text": "\u041e\u0442\u0432\u0435\u0442: B", "bold_text": "\u041e\u0442\u0432\u0435\u0442: B", "has_bold": True},
    ]

    items = parse_blocks_to_items(blocks)

    assert len(items) == 1
    assert items[0]["question"] == "Capital of Kazakhstan"
    assert items[0]["options"] == ["Almaty", "Astana", "Shymkent", "Aktobe"]
    assert items[0]["correct"] == 2
    assert items[0]["correct_answer"] == "Astana"
    assert items[0]["type"] == "multiple_choice"
    assert items[0]["distractors_source"] == "source_document_answer_line"


def test_parse_blocks_extracts_answer_prompt_followed_by_bold_options():
    blocks = [
        {"type": "text", "text": "Последствие перехода к НЭП:"},
        {"type": "text", "text": "A) КазАССР заняла 7-е место по валовой продукции"},
        {"type": "text", "text": "B) голод и потеря основной силы производства"},
        {"type": "text", "text": "C) налажено передовое производство цветных металлов и никеля"},
        {
            "type": "text",
            "text": "D) заработали сотни заводов и железнодорожные перевозки",
            "bold_text": "D) заработали сотни заводов и железнодорожные перевозки",
            "has_bold": True,
        },
    ]

    items = parse_blocks_to_items(blocks)

    assert len(items) == 1
    assert items[0]["question"] == "Последствие перехода к НЭП"
    assert items[0]["correct"] == 4
    assert items[0]["correct_answer"] == "заработали сотни заводов и железнодорожные перевозки"
    assert items[0]["options"] == [
        "КазАССР заняла 7-е место по валовой продукции",
        "голод и потеря основной силы производства",
        "налажено передовое производство цветных металлов и никеля",
        "заработали сотни заводов и железнодорожные перевозки",
    ]


def test_parse_blocks_splits_inline_options_and_uses_bold_answer_text():
    blocks = [
        {"type": "text", "text": "Город, взятый казахами:"},
        {
            "type": "text",
            "text": "A) Хивы B) Коканда C) Бухары D) Ферганы",
            "bold_text": "Хивы",
            "has_bold": True,
        },
    ]

    items = parse_blocks_to_items(blocks)

    assert len(items) == 1
    assert items[0]["options"] == ["Хивы", "Коканда", "Бухары", "Ферганы"]
    assert items[0]["correct"] == 1
    assert items[0]["correct_answer"] == "Хивы"


def test_parse_blocks_extracts_ready_quiz_with_multiple_bold_correct_options():
    blocks = [
        {"type": "text", "text": "Выберите верные утверждения:"},
        {"type": "text", "text": "A) Был создан общеказахский ополчение", "bold_text": "A) Был создан общеказахский ополчение", "has_bold": True},
        {"type": "text", "text": "B) Был заключен мир с джунгарами"},
        {"type": "text", "text": "C) Объединились силы трех жузов", "bold_text": "C) Объединились силы трех жузов", "has_bold": True},
        {"type": "text", "text": "D) Столица перенесена в Туркестан"},
    ]

    items = parse_blocks_to_items(blocks)

    assert len(items) == 1
    assert items[0]["correct"] == [1, 3]
    assert items[0]["correct_answers"] == [
        "Был создан общеказахский ополчение",
        "Объединились силы трех жузов",
    ]
    assert items[0]["type"] == "multiple_answer"
    assert items[0]["distractors_source"] == "source_document_bold"


def test_parse_blocks_splits_adjacent_option_groups_when_next_question_has_no_colon():
    blocks = [
        {"type": "text", "text": "Одно из достижений индустриализации в Казахстане:"},
        {"type": "text", "text": "A) добыча черных металлов"},
        {"type": "text", "text": "B) производство стройматериалов"},
        {"type": "text", "text": "C) производство энергии"},
        {"type": "text", "text": "D) добыча цветных металлов", "bold_text": "D) добыча цветных металлов", "has_bold": True},
        {"type": "text", "text": "Последствие насильственной коллективизации"},
        {"type": "text", "text": "A) развитие тяжелой промышленности"},
        {"type": "text", "text": "B) переход к оседлости", "bold_text": "B) переход к оседлости", "has_bold": True},
        {"type": "text", "text": "C) укрепление индивидуальных хозяйств"},
        {"type": "text", "text": "D) распространение кочевничества"},
    ]

    items = parse_blocks_to_items(blocks)

    assert len(items) == 2
    assert items[0]["question"] == "Одно из достижений индустриализации в Казахстане"
    assert items[0]["correct"] == 4
    assert items[1]["question"] == "Последствие насильственной коллективизации"
    assert items[1]["correct"] == 2


def test_parse_blocks_does_not_treat_abbreviated_name_as_option_marker():
    blocks = [
        {"type": "text", "text": "Хан, спасший А.Тевкелева от гибели:"},
        {"type": "text", "text": "А) Абылай"},
        {"type": "text", "text": "В) Абильмансур"},
        {"type": "text", "text": "С) Абулмамбет"},
        {"type": "text", "text": "Д) Абулхаир", "bold_text": "Д) Абулхаир", "has_bold": True},
        {"type": "text", "text": "А.Тевкелев был отправлен делегацией в Младший жуз по поручению:"},
        {"type": "text", "text": "А) Екатерины Великой"},
        {"type": "text", "text": "В) Анны Иоанновны", "bold_text": "В) Анны Иоанновны", "has_bold": True},
        {"type": "text", "text": "С) Елизаветы 2"},
        {"type": "text", "text": "Д) Анны Леопольдовны"},
    ]

    items = parse_blocks_to_items(blocks)

    assert len(items) == 2
    assert items[0]["correct"] == 4
    assert items[0]["correct_answer"] == "Абулхаир"
    assert items[1]["question"] == "А.Тевкелев был отправлен делегацией в Младший жуз по поручению"
    assert items[1]["correct"] == 2
    assert items[1]["correct_answer"] == "Анны Иоанновны"


def test_parse_blocks_extracts_numbered_statement_multi_answer():
    blocks = [
        {"type": "text", "text": "20. Установите верные утверждения:"},
        {"type": "text", "text": "1. Это делалось учитывая интересы депортируемой нации"},
        {"type": "text", "text": "2. Идея о создании немецкой автономии было инициативой казахов"},
        {"type": "text", "text": "3. Это подняло бы авторитет Казахстана"},
        {"type": "text", "text": "4. Немецкая автономия должна подчиняться республиканским органам"},
        {"type": "text", "text": "1 и 4", "bold_text": "1 и 4", "has_bold": True},
    ]

    items = parse_blocks_to_items(blocks)

    assert len(items) == 1
    assert items[0]["question"] == "Установите верные утверждения"
    assert items[0]["options"] == [
        "Это делалось учитывая интересы депортируемой нации",
        "Идея о создании немецкой автономии было инициативой казахов",
        "Это подняло бы авторитет Казахстана",
        "Немецкая автономия должна подчиняться республиканским органам",
    ]
    assert items[0]["correct"] == [1, 4]
    assert items[0]["correct_answers"] == [
        "Это делалось учитывая интересы депортируемой нации",
        "Немецкая автономия должна подчиняться республиканским органам",
    ]
    assert items[0]["distractors_source"] == "source_document_answer_indexes"


def test_parse_blocks_keeps_multiline_quote_as_context_not_question():
    blocks = [
        {"type": "text", "text": "19 мая"},
        {"type": "text", "text": "ОБЕД"},
        {"type": "text", "text": "10."},
        {
            "type": "text",
            "text": "“Среди политических соперников Есим-хана источники выделяют султана Турсуна.",
        },
        {
            "type": "text",
            "text": "Стремление Турсуна к власти, непризнание им хана Есим привело к двоевластию.",
        },
        {
            "type": "text",
            "text": "Но в 1627 году Турсун напал на его ставку - Туркестан. Текст про Турсуна.”",
        },
        {"type": "text", "text": "11. Хан, правивший в 1598 году: Есим"},
        {"type": "text", "text": "12. Система, при которой два человека стоят у власти - двоевластие"},
        {"type": "text", "text": "«Настанет ли день,"},
        {"type": "text", "text": "Когда нам удастся сесть"},
        {"type": "text", "text": "На рыжих, звонко ржащих коней!"},
        {"type": "text", "text": "Удастся ли нам преследовать бегущего врага?»."},
        {"type": "text", "text": "16. Тема его стихотворения: национальная независимость"},
        {"type": "text", "text": "17. Жырау того времени писали: толгау"},
    ]

    items = parse_blocks_to_items(blocks)

    assert [item["correct_answer"] for item in items] == [
        "Туркестан",
        "Есим",
        "двоевластие",
        "национальная независимость",
        "толгау",
    ]
    assert "Когда нам удастся сесть" not in [item["question"] for item in items]
    assert all("Турсуна" in item["context"] for item in items[:3])
    assert all("Когда нам удастся сесть" in item["context"] for item in items[3:])


def test_oversize_bold_correct_answer_is_flagged_not_crashing_other_questions():
    long_answer = "О" * (OPTION_MAX_CHARS + 20)
    blocks = [
        {"type": "text", "text": "Слишком длинный готовый вопрос:"},
        {"type": "text", "text": "A) короткий неверный"},
        {
            "type": "text",
            "text": f"B) {long_answer}",
            "bold_text": f"B) {long_answer}",
            "has_bold": True,
        },
        {"type": "text", "text": "Кто считается основателем Казахского ханства?"},
        {"type": "text", "text": "A) Абылай хан"},
        {
            "type": "text",
            "text": "B) Керей и Жанибек",
            "bold_text": "Керей и Жанибек",
            "has_bold": True,
        },
        {"type": "text", "text": "C) Кенесары хан"},
        {"type": "text", "text": "D) Тауке хан"},
    ]

    items = parse_blocks_to_items(blocks)

    # Один oversize вопрос помечен на ревью и не уронил парсинг остальных.
    assert len(items) == 2
    assert items[0]["type"] == "needs_option_length_review"
    assert items[0]["correct_answer"] == long_answer  # текст не обрезан, не потерян
    # Следующий нормальный вопрос распарсен корректно.
    assert items[1]["type"] == "multiple_choice"
    assert items[1]["correct_answer"] == "Керей и Жанибек"


def test_attach_distractors_flags_oversize_correct_answer_instead_of_building_option():
    long_answer = "Я" * (OPTION_MAX_CHARS + 5)
    items = attach_distractors(
        [
            {
                "question": "Локальный вопрос про хана",
                "correct_answer": long_answer,
                "options": [],
            }
        ]
    )

    assert len(items) == 1
    # attach_distractors не должен делать oversize-ответ вариантом опроса —
    # текст не теряется, вопрос помечается на ревью.
    assert items[0]["correct_answer"] == long_answer
    assert items[0]["type"] == "needs_option_length_review"
    assert items[0]["options"] == []
    assert items[0]["correct"] is None
    assert items[0]["distractors_source"] == "needs_option_length_review"


def test_count_unsupported_media_detects_image_inside_table(tmp_path):
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell_paragraph = table.cell(0, 0).paragraphs[0]
    run = cell_paragraph.add_run()
    image_path = tmp_path / "pic.png"
    # Minimal 1x1 PNG bytes.
    image_path.write_bytes(
        bytes.fromhex(
            "89504e470d0a1a0a0000000d49484452000000010000000108020000009077"
            "53de0000000c49444154789c63f8cfc0000003010100c9fe92ef0000000049"
            "454e44ae426082"
        )
    )
    run.add_picture(str(image_path))

    assert count_unsupported_media(doc) == 1


def test_build_output_reports_unsupported_table_media_warning(tmp_path):
    doc = Document()
    doc.add_paragraph("11 мая")
    doc.add_paragraph("Кто считается основателем Казахского ханства?")
    doc.add_paragraph("Керей и Жанибек")
    table = doc.add_table(rows=1, cols=1)
    run = table.cell(0, 0).paragraphs[0].add_run()
    image_path = tmp_path / "pic.png"
    image_path.write_bytes(
        bytes.fromhex(
            "89504e470d0a1a0a0000000d49484452000000010000000108020000009077"
            "53de0000000c49444154789c63f8cfc0000003010100c9fe92ef0000000049"
            "454e44ae426082"
        )
    )
    run.add_picture(str(image_path))

    docx_path = tmp_path / "with_table_media.docx"
    doc.save(docx_path)

    data = build_output(
        docx_path,
        tmp_path / "out.json",
        tmp_path / "media",
        "T",
        "D",
    )

    assert data["report"]["unsupported_media_locations"] == 1
    assert "warnings" in data["report"]
    assert "таблиц" in data["report"]["warnings"][0]
