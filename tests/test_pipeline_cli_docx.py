import json

from docx import Document

from backend import pipeline_cli


def make_docx(path, *, conflict=False):
    doc = Document()
    doc.add_paragraph("Тимур", style="Heading 1")
    doc.add_paragraph("На портрете изображен:")
    doc.add_paragraph("A) Абылай")
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("B) Эмир Тимур")
    run.bold = True
    doc.add_paragraph("C) Керей")
    if conflict:
        doc.add_paragraph("Ответ: A")
    doc.save(path)


def test_prepare_docx_cli_writes_artifacts_and_russian_summary(tmp_path, capsys):
    docx_path = tmp_path / "source.docx"
    out_dir = tmp_path / "out"
    make_docx(docx_path)

    exit_code = pipeline_cli.run([
        "prepare-docx",
        "--file",
        str(docx_path),
        "--strategy",
        "google-docs-docx-prep",
        "--out",
        str(out_dir),
    ])

    captured = capsys.readouterr()
    assert exit_code == 0
    assert "Подготовка DOCX завершена" in captured.out
    assert (out_dir / "prepared.md").exists()
    assert (out_dir / "prepared.json").exists()
    assert (out_dir / "report.json").exists()
    prepared = (out_dir / "prepared.md").read_text(encoding="utf-8")
    assert "#SECTION: Тимур" in prepared
    assert "#A*\nЭмир Тимур" in prepared
    assert json.loads((out_dir / "prepared.json").read_text(encoding="utf-8"))["question_count"] == 1


def test_prepare_docx_cli_returns_nonzero_but_writes_report_on_hard_error(tmp_path, capsys):
    docx_path = tmp_path / "source.docx"
    out_dir = tmp_path / "out"
    make_docx(docx_path, conflict=True)

    exit_code = pipeline_cli.run(["prepare-docx", "--file", str(docx_path), "--out", str(out_dir)])

    captured = capsys.readouterr()
    assert exit_code == 1
    assert "Найдены блокирующие ошибки" in captured.err
    report = json.loads((out_dir / "report.json").read_text(encoding="utf-8"))
    assert report["broken_questions"][0]["issue_code"] == "answer_bold_and_line_conflict"


def test_parse_prepared_cli_writes_clean_audit_and_report(tmp_path, capsys):
    prepared_path = tmp_path / "prepared.md"
    prepared_path.write_text(
        """#SECTION: Тимур
#CONTEXT
![media](media/timur.png)
Текст контекста
#END_CONTEXT
#Q
На портрете изображен:
#A
Абылай
#A*
Эмир Тимур
#END_Q
""",
        encoding="utf-8",
    )
    out_dir = tmp_path / "parsed"

    exit_code = pipeline_cli.run([
        "parse-prepared",
        "--file",
        str(prepared_path),
        "--strategy",
        "docx-strict-template",
        "--out",
        str(out_dir),
        "--title",
        "История Казахстана",
    ])

    captured = capsys.readouterr()
    assert exit_code == 0
    assert "Строгий парсинг завершен" in captured.out
    clean = json.loads((out_dir / "quiz.clean.json").read_text(encoding="utf-8"))
    audit = json.loads((out_dir / "quiz.audit.json").read_text(encoding="utf-8"))
    report = json.loads((out_dir / "report.json").read_text(encoding="utf-8"))
    assert clean["title"] == "История Казахстана"
    assert [item["type"] for item in clean["items"]] == ["title", "context", "question"]
    assert clean["items"][2]["answers"] == [2]
    assert audit["question_count"] == 1
    assert report["requires_review"] is False


def test_parse_prepared_cli_returns_nonzero_and_report_on_hard_error(tmp_path, capsys):
    prepared_path = tmp_path / "broken.md"
    prepared_path.write_text(
        """#Q
Вопрос
#A
Один
#A
Два
#END_Q
""",
        encoding="utf-8",
    )
    out_dir = tmp_path / "parsed"

    exit_code = pipeline_cli.run(["parse-prepared", "--file", str(prepared_path), "--out", str(out_dir)])

    captured = capsys.readouterr()
    assert exit_code == 1
    assert "Найдены блокирующие ошибки" in captured.err
    report = json.loads((out_dir / "report.json").read_text(encoding="utf-8"))
    assert report["error_count"] >= 1
    assert report["issues"][0]["message_ru"]


def test_pipeline_cli_help_does_not_import_upload_stack(capsys):
    parser = pipeline_cli.build_parser()
    help_text = parser.format_help()

    assert "prepare-docx" in help_text
    assert "parse-prepared" in help_text
    assert "Telegram" not in help_text
    assert "OpenAI" not in help_text