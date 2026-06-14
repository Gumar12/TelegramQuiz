# -*- coding: utf-8 -*-
"""Re-split the source quiz DOCX into N parts with a target number of parser-questions
each, cutting only at real question/context boundaries so no question is broken and
no phantom items are created. Formatting and images are preserved.

Counting authority: backend.docx_to_quiz_json_v2.parse_blocks_to_items (556 items).
"""
from __future__ import annotations
import copy
import sys
from pathlib import Path

sys.path.insert(0, r"C:\Users\Asus\Documents\Agentic\Quizbot")

from docx import Document
from docx.oxml.ns import qn

from backend.docx_to_quiz_json_v2 import (
    clean, run_is_bold, parse_blocks_to_items,
    split_option_line, NUMBERED_RE, ANSWER_CHOICE_RE,
)
import re

ROMAN_OPT_RE = re.compile(r"^\s*(?:I{1,3}|IV|V|VI{0,3}|IX|X)[\.\)]\s")
QUOTE_OPENERS = ("«", '"', "“")


def classify(text):
    """'opt' = option/continuation line that cannot begin a part;
    'start' = a question stem or context title/quote that can begin a part."""
    t = clean(text)
    if not t:
        return "empty"
    if split_option_line(t):
        return "opt"
    if NUMBERED_RE.match(t):
        return "opt"
    if ROMAN_OPT_RE.match(t):
        return "opt"
    if ANSWER_CHOICE_RE.match(t):
        return "opt"
    if re.match(r"^\s*ответ\b", t, re.I):
        return "opt"
    return "start"

SRC = Path(r"C:\Users\Asus\Downloads\мне надо-2.docx")
OUT_DIR = SRC.parent
TARGETS = [200, 200, None]          # last None = "the rest"
OUT_NAMES = [
    "мне надо-2_часть1_1-200.docx",
    "мне надо-2_часть2_201-400.docx",
    "мне надо-2_часть3_401-556.docx",
]


def paragraph_blocks(paras):
    """Return (para_index, block) list mirroring iter_docx_blocks, but in-memory
    and WITHOUT writing image blobs (dummy paths — only types/structure matter)."""
    out = []
    for pi, para in enumerate(paras):
        text_parts, bold_parts, n_imgs = [], [], 0
        for run in para.runs:
            if run.text:
                text_parts.append(run.text)
                if run_is_bold(run):
                    bold_parts.append(run.text)
            for drawing in run._element.xpath(".//w:drawing"):
                for blip in drawing.xpath(".//a:blip"):
                    if blip.get(qn("r:embed")):
                        n_imgs += 1
        text = clean("".join(text_parts))
        if text:
            block = {"type": "text", "text": text}
            bold = clean("".join(bold_parts))
            if bold:
                block["bold_text"] = bold
                block["has_bold"] = True
            out.append((pi, block))
        for _ in range(n_imgs):
            out.append((pi, {"type": "image", "path": f"img_{pi}"}))
    return out


def count_prefix(all_blocks, k):
    """Number of parser-items in paragraphs [0:k)."""
    blocks = [b for pi, b in all_blocks if pi < k]
    return len(parse_blocks_to_items(blocks))


def count_range(all_blocks, a, b):
    blocks = [blk for pi, blk in all_blocks if a <= pi < b]
    return len(parse_blocks_to_items(blocks))


def clean_boundaries(paras):
    """Paragraph indices that safely begin a new question/context unit:
    a 'start' line whose previous non-empty line was an option ('opt') —
    i.e. the first content after a completed question's options."""
    bounds = []
    prev_class = None
    for i, para in enumerate(paras):
        c = classify(para.text)
        if c == "empty":
            continue
        if c == "start" and prev_class == "opt":
            bounds.append(i)
        prev_class = c
    return bounds


def choose_cut(all_blocks, bounds, target, after):
    """Among boundary indices > `after`, pick the one whose prefix has exactly
    `target` complete items; fall back to the closest if no exact match."""
    candidates = [b for b in bounds if b > after]
    best = None
    for b in candidates:
        c = count_prefix(all_blocks, b)
        if c == target:
            return b, c          # first exact match = cut right at item target+1's start
        if c > target:
            # overshot; return closest of this and the previous candidate
            return (best if best else b), (count_prefix(all_blocks, best) if best else c)
        best = b
    return (best if best else after), (count_prefix(all_blocks, best) if best else 0)


def prune_orphan_images(doc):
    """Drop image relationships no longer referenced by any kept paragraph,
    so removed-paragraph images don't bloat the output package."""
    referenced = set()
    for blip in doc.element.body.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid:
            referenced.add(rid)
    dropped = 0
    for rid, rel in list(doc.part.rels.items()):
        if rel.reltype.endswith("/image") and rid not in referenced:
            doc.part.drop_rel(rid)
            dropped += 1
    return dropped


def write_slice(src_path, out_path, a, b, total_paras):
    """Clone the source doc and keep only body paragraphs with index in [a,b)."""
    doc = Document(src_path)
    paras = doc.paragraphs
    assert len(paras) == total_paras, "paragraph count drifted"
    for idx, para in enumerate(paras):
        if not (a <= idx < b):
            para._p.getparent().remove(para._p)
    dropped = prune_orphan_images(doc)
    doc.save(out_path)
    return dropped


def main():
    doc = Document(SRC)
    paras = doc.paragraphs
    total_paras = len(paras)
    all_blocks = paragraph_blocks(paras)
    grand_total = len(parse_blocks_to_items([b for _, b in all_blocks]))
    print(f"source paragraphs: {total_paras}")
    print(f"source parser-items: {grand_total}")

    bounds = clean_boundaries(paras)

    # Find cut points (paragraph indices) for cumulative targets.
    cuts = [0]
    cumulative = 0
    for t in TARGETS[:-1]:
        cumulative += t
        k, got = choose_cut(all_blocks, bounds, cumulative, cuts[-1])
        if got != cumulative:
            print(f"  NOTE: nearest clean boundary gives {got} items (target {cumulative})")
        cuts.append(k)
    cuts.append(total_paras)

    # Report and write.
    print(f"cut paragraph indices: {cuts}")
    for i in range(len(OUT_NAMES)):
        a, b = cuts[i], cuts[i + 1]
        n = count_range(all_blocks, a, b)
        out_path = OUT_DIR / OUT_NAMES[i]
        dropped = write_slice(SRC, out_path, a, b, total_paras)
        kept_imgs = sum(1 for pi, blk in all_blocks if a <= pi < b and blk["type"] == "image")
        # boundary sanity: first non-empty text of the slice
        head = next((blk["text"] for pi, blk in all_blocks
                     if a <= pi and blk["type"] == "text"), "")
        print(f"  {OUT_NAMES[i]}: paras[{a}:{b}]  items={n}  images_kept={kept_imgs} "
              f"orphans_dropped={dropped}  head={head[:50]!r}")

    # Dump the text around each cut (decoded UTF-8) for context-integrity review.
    insp = OUT_DIR.parent / "boundary_review.txt" if False else \
        Path(r"C:\Users\Asus\Documents\Agentic\Quizbot\scripts\_boundary_review.txt")
    nonempty = [(pi, blk["text"]) for pi, blk in all_blocks if blk["type"] == "text"]
    with open(insp, "w", encoding="utf-8") as f:
        for cut in cuts[1:-1]:
            f.write(f"================ CUT at paragraph {cut} ================\n")
            f.write("---- tail of previous part ----\n")
            for pi, t in nonempty:
                if cut - 12 <= pi < cut:
                    f.write(f"[{pi}] {t}\n")
            f.write("---- head of next part (starts here) ----\n")
            for pi, t in nonempty:
                if cut <= pi < cut + 12:
                    f.write(f"[{pi}] {t}\n")
            f.write("\n")
    print(f"boundary review written to {insp}")

    print("\nVerifying written files with the project parser...")
    import tempfile
    from backend.docx_to_quiz_json_v2 import iter_docx_blocks
    total_check = 0
    for name in OUT_NAMES:
        p = OUT_DIR / name
        items = parse_blocks_to_items(iter_docx_blocks(p, tempfile.mkdtemp()))
        total_check += len(items)
        print(f"  {name}: {len(items)} items")
    print(f"  sum = {total_check} (source = {grand_total})")


if __name__ == "__main__":
    main()
